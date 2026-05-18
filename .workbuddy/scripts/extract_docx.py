#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
import sys
import io
from pathlib import Path
from docx import Document

WORKSPACE = Path(__file__).resolve().parent.parent.parent  # 从 .workbuddy/scripts/ 向上三级


def extract_docx(file_path):
    """Extract text from docx file"""
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    text.append(" | ".join(row_text))

        return "\n".join(text[:100])
    except Exception as e:
        return f"Error: {e}"


def main():
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    docx_files = [
        "05-知识沉淀/平台运营制度/平台制度V3.2.docx",
        "05-知识沉淀/结算与提成体系/平台结算制度V3.2.docx",
        "05-知识沉淀/师傅培训体系/师傅运营核心.docx",
        "05-知识沉淀/师傅培训体系/SOP-3.2.2.docx",
    ]

    print("Extracting docx content...")

    for f in docx_files:
        path = WORKSPACE / f
        if path.exists():
            name = path.stem
            print(f"Processing: {name}")
            content = extract_docx(path)
            output = WORKSPACE / "05-知识沉淀" / "知识结晶" / f"{name}_提取.md"
            with open(output, "w", encoding="utf-8") as wf:
                wf.write(
                    f"# {name}\n\n> Extract time: 2026-04-10\n\n---\n\n{content[:3000]}"
                )
            print(f"Saved: {output.name}")
        else:
            print(f"Not found: {f}")


if __name__ == "__main__":
    main()
